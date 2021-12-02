from copy import deepcopy
from pathlib import Path
from win32com import client as wc  # pip install pypiwin32
from docx import Document  # pip install python-docx
import pandas as pd


# python-docx不能处理doc文档，使用win32com转存为docx文档
def doctransform2docx(doc_path):
    docx_path = doc_path + 'x'
    suffix = doc_path.split('.')[1]
    assert 'doc' in suffix, '传入的不是word文档，请重新输入！'
    if suffix == 'docx':
        return Document(doc_path)
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(doc_path)
    doc.SaveAs2(docx_path, 16)  # docx为16
    doc.Close()
    word.Quit()
    return Document(docx_path)


# 替换docx中的特定字符，由于run方法在有格式的docx文件中展示效果很差，故将docx中的文本的需要填充出英文字符占位
def replace_docx(name, values, wordfile, path_name='Company'):
    wordfile_copy = deepcopy(wordfile)  # 防止原文件被篡改，deepcopy为副本
    for col_name, value in zip(name, values):
        if col_name == 'Company':
            path_name = str(value)
        for paragraphs in wordfile_copy.paragraphs:
            for run in paragraphs.runs:
                run.text = run.text.replace(col_name, str(value))
    # docx文档替换完毕，另存为，一定要用绝对路径
    wordfile_copy.save(f'{save_folder}/{path_name}.docx')


if __name__ == '__main__':
    # 定义需处理的文件路径
    doc_path = r"D:\solve_path\单位.doc"
    excel_path = r"D:\solve_path\信息.xls"
    save_folder = Path('D:/docx_save')
    save_folder.mkdir(parents=True, exist_ok=True)  # 文件夹没有时自动创建
    # 获取excel数据
    data = pd.read_excel(excel_path)
    wordfile = doctransform2docx(doc_path)
    data_save = data.apply(lambda x: replace_docx(x.index, x.values, wordfile), axis=1)
